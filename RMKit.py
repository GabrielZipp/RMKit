# ============================================
# RMKIT - KIT COMPLETO EM PDF (VIA UNIÃO DE PDFs)
# ============================================

import os
import tempfile
from datetime import datetime
from docx import Document
import comtypes.client
import PyPDF2  

class GeradorKitPDF:
    def __init__(self):
        self.temp_dir = tempfile.mkdtemp()
        
    def data_por_extenso(self):
        meses = {1:'janeiro',2:'fevereiro',3:'março',4:'abril',5:'maio',6:'junho',
                7:'julho',8:'agosto',9:'setembro',10:'outubro',11:'novembro',12:'dezembro'}
        hoje = datetime.now()
        return f"{hoje.day} de {meses[hoje.month]} de {hoje.year}"
    
    def substituir_no_documento(self, doc, dados):
        """Substitui placeholders no documento"""
        for paragraph in doc.paragraphs:
            for var, valor in dados.items():
                if var in paragraph.text:
                    for run in paragraph.runs:
                        if var in run.text:
                            run.text = run.text.replace(var, str(valor))
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for var, valor in dados.items():
                            if var in paragraph.text:
                                for run in paragraph.runs:
                                    if var in run.text:
                                        run.text = run.text.replace(var, str(valor))
        return doc
    
    def gerar_pdf_individual(self, template_path, dados, output_pdf):
        """
        Gera um PDF a partir de um template DOCX
        Retorna o caminho do PDF gerado
        """
        # Carregar template
        doc = Document(template_path)
        doc = self.substituir_no_documento(doc, dados)
        
        # Salvar DOCX temporário
        temp_docx = os.path.join(self.temp_dir, os.path.basename(template_path).replace('.docx', '_temp.docx'))
        doc.save(temp_docx)
        
        # Converter para PDF usando Word
        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False
        doc_word = word.Documents.Open(os.path.abspath(temp_docx))
        doc_word.SaveAs(os.path.abspath(output_pdf), FileFormat=17)
        doc_word.Close()
        word.Quit()
        
        return output_pdf
    
    def gerar_kit_unico(self, dados_cliente, nome_saida=None):
        """
        Gera procuração e contrato em PDF e depois os une em um único arquivo
        """
        # Adicionar data atual
        dados_cliente['{{DATA_EXTENSO}}'] = self.data_por_extenso()
        
        print(f"\n📄 Gerando kit para: {dados_cliente['{{NOME}}']}")
        print("-" * 60)
        
        # 1. Gerar PDF da procuração
        print("1️⃣  Gerando procuração PDF...")
        pdf_proc = os.path.join(self.temp_dir, "procuracao.pdf")
        self.gerar_pdf_individual("procuração.docx", dados_cliente.copy(), pdf_proc)
        print("   ✅ Procuração PDF criada")
        
        # 2. Gerar PDF do contrato
        print("2️⃣  Gerando contrato PDF...")
        pdf_cont = os.path.join(self.temp_dir, "contrato.pdf")
        self.gerar_pdf_individual("contrato.docx", dados_cliente.copy(), pdf_cont)
        print("   ✅ Contrato PDF criado")
        
        # 3. Unir os dois PDFs
        print("3️⃣  Unindo PDFs...")
        if nome_saida is None:
            nome_saida = f"Kit_{dados_cliente['{{NOME}}'].replace(' ', '_')}.pdf"
        
        merger = PyPDF2.PdfMerger()
        merger.append(pdf_proc)
        merger.append(pdf_cont)
        merger.write(nome_saida)
        merger.close()
        
        print(f"   ✅ PDF único gerado: {nome_saida}")
        
        # Limpeza (opcional)
        try:
            os.remove(pdf_proc)
            os.remove(pdf_cont)
            os.rmdir(self.temp_dir)
        except:
            pass
        
        print("\n" + "="*60)
        print(f"✨ KIT COMPLETO CRIADO: {os.path.abspath(nome_saida)}")
        print("="*60)
        
        return nome_saida


# ============================================
# EXECUÇÃO
# ============================================

if __name__ == "__main__":
    
    print("="*70)
    print("🚀 RMKIT - KIT PROCURAÇÃO + CONTRATO EM UM ÚNICO PDF")
    print("="*70)
    
    # ===== DADOS DO CLIENTE (EDITAR AQUI) =====
    dados_cliente = {
        '{{NOME}}': 'MARIA APARECIDA SANTOS',
        '{{NACIONALIDADE}}': 'brasileira',
        '{{ESTADO_CIVIL}}': 'solteira',
        '{{PROFISSAO}}': 'comerciante',
        '{{RG}}': '12.345.678-9 SSP/SP',
        '{{CPF}}': '123.456.789-00',
        '{{RUA}}': 'das Flores',
        '{{NUMERO}}': '123',
        '{{BAIRRO}}': 'Centro',
        '{{CEP}}': '13880-000',
        '{{CIDADE}}': 'Aguaí',
        '{{ESTADO}}': 'SP',
        '{{TIPO_ACAO}}': 'Trabalhista',
        '{{HONORARIOS_FIXOS}}': '03 (três) salários-mínimos',
        '{{HONORARIOS_EXITO}}': '30% (trinta por cento)',
        '{{MULTA_RESCISAO}}': '01 (um) salário-mínimo'
    }
    
    # Mostrar resumo
    print("\n📋 Dados do cliente:")
    for campo in ['{{NOME}}', '{{TIPO_ACAO}}', '{{HONORARIOS_FIXOS}}', '{{MULTA_RESCISAO}}']:
        print(f"   {campo}: {dados_cliente[campo]}")
    
    # Gerar kit
    gerador = GeradorKitPDF()
    gerador.gerar_kit_unico(dados_cliente)
