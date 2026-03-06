# ============================================
# RMKIT - TESTE SIMPLES (primeiro_test.py)
# ============================================

from docx import Document
from datetime import datetime
import os
import comtypes.client
import time

def testar_conversao_simples():
    """Teste básico para ver se a conversão funciona"""
    
    print("="*60)
    print("🔧 TESTE DE CONVERSÃO DOCX -> PDF")
    print("="*60)
    
    # 1. Criar um documento de teste bem simples
    print("\n1. Criando documento de teste...")
    doc = Document()
    doc.add_heading('TESTE DE CONVERSÃO', 0)
    doc.add_paragraph('Este é um documento de teste para verificar a conversão para PDF.')
    doc.add_paragraph('Data: {}'.format(datetime.now().strftime('%d/%m/%Y')))
    
    teste_docx = "teste_simples.docx"
    doc.save(teste_docx)
    print(f"   ✅ Documento criado: {teste_docx}")
    
    # 2. Tentar converter
    print("\n2. Tentando converter para PDF...")
    teste_pdf = "teste_simples.pdf"
    
    try:
        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False
        print("   ✅ Word iniciado")
        
        doc_word = word.Documents.Open(os.path.abspath(teste_docx))
        print("   ✅ Documento aberto no Word")
        
        doc_word.SaveAs(os.path.abspath(teste_pdf), FileFormat=17)
        print("   ✅ Documento salvo como PDF")
        
        doc_word.Close()
        word.Quit()
        
        print(f"\n✅ SUCESSO! PDF gerado: {teste_pdf}")
        return True
        
    except Exception as e:
        print(f"\n❌ ERRO: {e}")
        return False

def testar_seu_template():
    """Testa apenas seu template sem junção"""
    
    print("\n" + "="*60)
    print("📄 TESTANDO SEU TEMPLATE")
    print("="*60)
    
    # Verificar se o template existe
    if not os.path.exists("procuração.docx"):
        print("❌ Template 'procuração.docx' não encontrado!")
        return False
    
    print(f"\n1. Template encontrado: procuração.docx")
    print(f"   Tamanho: {os.path.getsize('procuração.docx')} bytes")
    
    # 1. Abrir e fazer uma cópia simples
    print("\n2. Abrindo template...")
    try:
        doc = Document("procuração.docx")
        print("   ✅ Template aberto com sucesso")
        
        # 2. Salvar uma cópia
        copia = "copia_teste.docx"
        doc.save(copia)
        print(f"   ✅ Cópia salva: {copia}")
        
        # 3. Tentar converter a cópia
        print("\n3. Tentando converter cópia para PDF...")
        pdf_teste = "teste_template.pdf"
        
        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False
        
        doc_word = word.Documents.Open(os.path.abspath(copia))
        doc_word.SaveAs(os.path.abspath(pdf_teste), FileFormat=17)
        doc_word.Close()
        word.Quit()
        
        print(f"\n✅ SUCESSO! PDF gerado: {pdf_teste}")
        
        # Limpar
        os.remove(copia)
        return True
        
    except Exception as e:
        print(f"\n❌ ERRO: {e}")
        return False

# ============================================
# VERSÃO ALTERNATIVA (sem Word)
# ============================================

def salvar_como_pdf_sem_word():
    """Salva como PDF usando print do Windows (alternativa)"""
    
    print("\n" + "="*60)
    print("🖨️ MÉTODO ALTERNATIVO - Imprimir como PDF")
    print("="*60)
    
    import win32print
    import win32api
    
    try:
        # Abrir o documento no Word e mandar imprimir como PDF
        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = True  # Deixa visível para você ver
        
        doc = word.Documents.Open(os.path.abspath("procuração.docx"))
        
        # Configurar para imprimir como PDF
        word.ActiveDocument.PrintOut(
            Background=False,
            OutputFileName=os.path.abspath("impressao.pdf"),
            PrintToFile=True
        )
        
        doc.Close()
        word.Quit()
        
        print("✅ PDF gerado via impressão")
        return True
        
    except Exception as e:
        print(f"❌ Erro: {e}")
        return False

# ============================================
# EXECUTAR TESTES
# ============================================

if __name__ == "__main__":
    
    print("🔍 DIAGNÓSTICO DO SISTEMA")
    print("-" * 60)
    
    # 1. Teste simples primeiro
    if testar_conversao_simples():
        print("\n✅ Conversão básica funcionou!")
    else:
        print("\n❌ Conversão básica falhou. Problema com o Word?")
    
    # 2. Testar seu template
    testar_seu_template()
    
    print("\n" + "="*60)
    print("📋 PRÓXIMOS PASSOS:")
    print("="*60)
    print("1. O teste simples funcionou?")
    print("2. O teste do seu template funcionou?")
    print("3. Me diga os resultados para eu ajustar o código!")