# ============================================
# RMKIT - MÓDULO DE LÓGICA (gerador.py) - CORRIGIDO
# ============================================

import os
import tempfile
import time
from datetime import datetime
from docx import Document
import comtypes.client
import PyPDF2

class GeradorKit:
    """Classe responsável por toda a lógica de geração de documentos"""
    
    def __init__(self):
        self.temp_dir = None
        
    def data_por_extenso(self):
        """Retorna data atual por extenso"""
        meses = {1:'janeiro',2:'fevereiro',3:'março',4:'abril',5:'maio',6:'junho',
                7:'julho',8:'agosto',9:'setembro',10:'outubro',11:'novembro',12:'dezembro'}
        hoje = datetime.now()
        return f"{hoje.day} de {meses[hoje.month]} de {hoje.year}"
    
    def substituir_placeholders(self, doc, dados):
        """Substitui placeholders em um documento Word"""
        for paragraph in doc.paragraphs:
            for var, valor in dados.items():
                if var in paragraph.text:
                    for run in paragraph.runs:
                        if var in run.text:
                            run.text = run.text.replace(var, str(valor))
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for var, valor in dados.items():
                            if var in paragraph.text:
                                for run in paragraph.runs:
                                    if var in run.text:
                                        run.text = run.text.replace(var, str(valor))
        return doc
    
    def docx_para_pdf(self, docx_path, pdf_path):
        """Converte DOCX para PDF usando Microsoft Word (VERSÃO CORRIGIDA)"""
        word = None
        doc = None
        
        try:
            # Garantir caminhos absolutos
            docx_path = os.path.abspath(docx_path)
            pdf_path = os.path.abspath(pdf_path)
            
            print(f"   Convertendo: {os.path.basename(docx_path)}")
            
            # Iniciar Word
            word = comtypes.client.CreateObject("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False  # Desabilitar alertas
            
            # Aguardar um pouco para o Word iniciar
            time.sleep(1)
            
            # Abrir documento
            doc = word.Documents.Open(docx_path)
            
            # Aguardar documento carregar
            time.sleep(1)
            
            # Salvar como PDF
            doc.SaveAs(pdf_path, FileFormat=17)
            
            # Aguardar salvamento
            time.sleep(1)
            
            # Fechar documento
            doc.Close()
            
            print(f"   ✅ PDF gerado: {os.path.basename(pdf_path)}")
            return True
            
        except Exception as e:
            print(f"   ❌ Erro detalhado: {e}")
            raise Exception(f"Erro na conversão para PDF: {e}")
            
        finally:
            # Garantir que tudo seja fechado
            try:
                if doc:
                    doc.Close()
            except:
                pass
            
            try:
                if word:
                    word.Quit()
            except:
                pass
    
    def gerar_pdf_individual(self, template_path, dados, output_pdf):
        """Gera um PDF a partir de um template DOCX"""
        try:
            # Abrir template
            doc = Document(template_path)
            
            # Substituir placeholders
            doc = self.substituir_placeholders(doc, dados)
            
            # Criar pasta temporária
            self.temp_dir = tempfile.mkdtemp()
            temp_docx = os.path.join(self.temp_dir, "temp.docx")
            doc.save(temp_docx)
            
            # Converter para PDF
            self.docx_para_pdf(temp_docx, output_pdf)
            
            # Limpar temporário
            if os.path.exists(temp_docx):
                os.remove(temp_docx)
            if os.path.exists(self.temp_dir):
                os.rmdir(self.temp_dir)
            
            return output_pdf
            
        except Exception as e:
            # Limpeza em caso de erro
            try:
                if os.path.exists(temp_docx):
                    os.remove(temp_docx)
                if os.path.exists(self.temp_dir):
                    os.rmdir(self.temp_dir)
            except:
                pass
            raise e
    
    def gerar_kit(self, template_proc, template_cont, dados_cliente, output_pdf):
        """
        Gera um único PDF com procuração + contrato
        """
        # Validar templates
        if not os.path.exists(template_proc):
            raise FileNotFoundError(f"Template de procuração não encontrado: {template_proc}")
        
        if not os.path.exists(template_cont):
            raise FileNotFoundError(f"Template de contrato não encontrado: {template_cont}")
        
        # Adicionar data atual
        dados_cliente['{{DATA_EXTENSO}}'] = self.data_por_extenso()
        
        # Criar pasta temporária
        self.temp_dir = tempfile.mkdtemp()
        
        pdf_proc = os.path.join(self.temp_dir, "procuracao.pdf")
        pdf_cont = os.path.join(self.temp_dir, "contrato.pdf")
        
        try:
            # Gerar PDF da procuração
            print("\n📄 Gerando procuração...")
            self.gerar_pdf_individual(template_proc, dados_cliente.copy(), pdf_proc)
            
            # Gerar PDF do contrato
            print("📄 Gerando contrato...")
            self.gerar_pdf_individual(template_cont, dados_cliente.copy(), pdf_cont)
            
            # Unir PDFs
            print("📎 Unindo PDFs...")
            merger = PyPDF2.PdfMerger()
            merger.append(pdf_proc)
            merger.append(pdf_cont)
            merger.write(output_pdf)
            merger.close()
            
            # Limpeza
            if os.path.exists(pdf_proc):
                os.remove(pdf_proc)
            if os.path.exists(pdf_cont):
                os.remove(pdf_cont)
            if os.path.exists(self.temp_dir):
                os.rmdir(self.temp_dir)
            
            print(f"✅ Kit gerado: {output_pdf}")
            return output_pdf
            
        except Exception as e:
            # Limpeza em caso de erro
            try:
                if os.path.exists(pdf_proc):
                    os.remove(pdf_proc)
                if os.path.exists(pdf_cont):
                    os.remove(pdf_cont)
                if os.path.exists(self.temp_dir):
                    os.rmdir(self.temp_dir)
            except:
                pass
            raise e
    
    def validar_dados(self, dados):
        """Valida se os campos obrigatórios estão preenchidos"""
        obrigatorios = ['{{NOME}}', '{{CPF}}', '{{RG}}']
        faltando = [campo for campo in obrigatorios if not dados.get(campo)]
        return len(faltando) == 0, faltando